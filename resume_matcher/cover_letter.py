"""Cover Letter Generator: creates professional cover letters using LLM."""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from resume_matcher.llm_client import get_llm_client
from resume_matcher.prompts import COVER_LETTER_SYSTEM, COVER_LETTER_USER

logger = logging.getLogger(__name__)


@dataclass
class CoverLetterResult:
    """Result of cover letter generation."""

    text: str = ""
    paragraphs: list[str] = field(default_factory=list)
    word_count: int = 0

    def to_dict(self) -> dict:
        return {
            "cover_letter_text": self.text,
            "paragraphs": self.paragraphs,
            "word_count": self.word_count,
        }


class CoverLetterGenerator:
    """Generates cover letters using LLM and renders them as .docx."""

    def generate(self, resume_data: dict, job_data: dict) -> CoverLetterResult:
        """Generate a cover letter for a specific job application.

        Args:
            resume_data: ATS-standard resume (from ATSTransformer or parser)
            job_data: Parsed job description

        Returns:
            CoverLetterResult with text, paragraphs, and word count.
        """
        client = get_llm_client()

        prompt = COVER_LETTER_USER.format(
            resume_json=json.dumps(resume_data, indent=2)[:6000],
            job_json=json.dumps(job_data, indent=2)[:4000],
        )

        data = client.complete_json(COVER_LETTER_SYSTEM, prompt)

        result = CoverLetterResult()
        result.text = data.get("cover_letter_text", "")
        result.paragraphs = data.get("paragraphs", [])
        result.word_count = data.get("word_count", len(result.text.split()))

        if not result.text and result.paragraphs:
            result.text = "\n\n".join(result.paragraphs)

        return result

    def generate_docx(
        self,
        cover_letter: CoverLetterResult,
        candidate_name: str = "",
        company_name: str = "",
        job_title: str = "",
    ) -> bytes:
        """Render a cover letter as a formatted .docx file.

        Args:
            cover_letter: CoverLetterResult from generate()
            candidate_name: Candidate's full name (for header/closing)
            company_name: Target company name
            job_title: Target job title

        Returns:
            Bytes of the .docx file.
        """
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(6)

        # Candidate name header
        if candidate_name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = name_para.add_run(candidate_name)
            run.bold = True
            run.font.size = Pt(14)

        # Date
        from datetime import date

        date_para = doc.add_paragraph()
        date_para.add_run(date.today().strftime("%B %d, %Y"))

        # Greeting
        if company_name:
            doc.add_paragraph(f"Dear {company_name} Hiring Team,")
        else:
            doc.add_paragraph("Dear Hiring Manager,")

        # Body paragraphs
        paragraphs = cover_letter.paragraphs or [cover_letter.text]
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.space_after = Pt(8)

        # Closing
        doc.add_paragraph("Sincerely,")
        if candidate_name:
            closing = doc.add_paragraph()
            run = closing.add_run(candidate_name)
            run.bold = True

        # Write to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
