"""Resume Generator: creates a formatted .docx file from updated resume data."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


# Section rendering order
SECTION_ORDER = [
    "header",
    "summary",
    "skills",
    "experience",
    "education",
    "certifications",
    "projects",
    "awards",
    "languages",
    "references",
]


class ResumeGenerator:
    """Generates a professional .docx resume from structured data."""

    def generate(self, resume_data: dict) -> bytes:
        """Generate a .docx file from resume sections dict.

        Args:
            resume_data: Dict with "sections" key containing resume content.
                Each section has "heading" and "content" keys.

        Returns:
            Bytes of the generated .docx file.
        """
        sections = resume_data.get("sections", {})
        if not sections and "updated_sections" in resume_data:
            sections = resume_data["updated_sections"]

        doc = Document()
        self._set_document_style(doc)

        # Render sections in standard order
        for category in SECTION_ORDER:
            if category in sections:
                data = sections[category]
                if category == "header":
                    self._render_header(doc, data)
                else:
                    self._render_section(doc, category, data)

        # Render any remaining sections not in the standard order
        for category, data in sections.items():
            if category not in SECTION_ORDER:
                self._render_section(doc, category, data)

        # Write to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _set_document_style(self, doc: Document) -> None:
        """Set default document margins and font."""
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Reduce paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(2)

    def _render_header(self, doc: Document, data: dict) -> None:
        """Render the contact/name header section."""
        content = data.get("content", [])
        if not content:
            return

        # First line is typically the name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(content[0])
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Remaining lines are contact details
        for line in content[1:]:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(line)
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Add a thin line after header
        self._add_horizontal_rule(doc)

    def _render_section(
        self, doc: Document, category: str, data: dict
    ) -> None:
        """Render a standard resume section with heading and content."""
        heading = data.get("heading", category.title())
        content = data.get("content", [])

        if not content:
            return

        # Section heading
        heading_para = doc.add_paragraph()
        heading_para.paragraph_format.space_before = Pt(8)
        heading_para.paragraph_format.space_after = Pt(2)
        heading_run = heading_para.add_run(heading.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Add thin line under heading
        self._add_horizontal_rule(doc)

        # Content
        if category == "experience":
            self._render_experience(doc, content, entries=data.get("entries"))
        elif category == "skills":
            self._render_skills(doc, content)
        else:
            for line in content:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(1)
                run = para.add_run(line)
                run.font.size = Pt(10)

    def _render_experience(self, doc: Document, content: list[str], entries: list[dict] | None = None) -> None:
        """Render experience section with structured entries or flat content fallback."""
        if entries:
            self._render_experience_entries(doc, entries)
            return

        # Fallback: flat content rendering (legacy data without entries)
        import re

        for line in content:
            is_title = bool(
                re.match(r"^[A-Z][\w\s]+\s*[—\-–]\s*\w", line)
            )
            if is_title:
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(1)
                run = para.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
            else:
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.add_run(line)
                run.font.size = Pt(10)

    def _render_experience_entries(self, doc: Document, entries: list[dict]) -> None:
        """Render structured experience entries with company, title, dates, and bullets."""
        for entry in entries:
            company = entry.get("company", "")
            title = entry.get("title", "")
            location = entry.get("location", "")
            dates = entry.get("dates", "")
            bullets = entry.get("bullets", [])

            # Company + Location line
            if company:
                company_line = company
                if location:
                    company_line += f"  |  {location}"
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(0)
                run = para.add_run(company_line)
                run.bold = True
                run.font.size = Pt(10)

            # Title + Dates line
            if title or dates:
                title_line = title if title else ""
                if dates:
                    title_line += f"  |  {dates}" if title_line else dates
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(1)
                run = para.add_run(title_line)
                run.italic = True
                run.font.size = Pt(10)

            # Bullets
            for bullet in bullets:
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.add_run(bullet)
                run.font.size = Pt(10)

    def _render_skills(self, doc: Document, content: list[str]) -> None:
        """Render skills section."""
        for line in content:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run(line)
            run.font.size = Pt(10)

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Add a thin horizontal line."""
        from docx.oxml.ns import qn
        from lxml import etree

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        pPr = para._element.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
